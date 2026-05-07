"""
Professional Resume Builder with Premium Templates
Inspired by NovaResume.com design patterns and best practices
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import traceback


class ProfessionalResumeBuilder:
    """Build professional resumes with premium templates and formatting"""
    
    def __init__(self):
        self.templates = {
            "Modern Pro": self.build_modern_pro_template,
            "Executive": self.build_executive_template,
            "Minimalist": self.build_minimalist_template,
            "Corporate": self.build_corporate_template,
            "Creative Tech": self.build_creative_tech_template
        }
        
        # Color schemes for different templates
        self.color_schemes = {
            "Modern Pro": {"primary": RGBColor(31, 78, 121), "secondary": RGBColor(79, 129, 189)},
            "Executive": {"primary": RGBColor(0, 0, 0), "secondary": RGBColor(89, 89, 89)},
            "Minimalist": {"primary": RGBColor(64, 64, 64), "secondary": RGBColor(128, 128, 128)},
            "Corporate": {"primary": RGBColor(0, 51, 102), "secondary": RGBColor(0, 102, 204)},
            "Creative Tech": {"primary": RGBColor(52, 152, 219), "secondary": RGBColor(155, 89, 182)}
        }

    def generate_resume(self, data, template_name="Modern Pro"):
        """Generate professional resume based on template"""
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.6)
                section.bottom_margin = Inches(0.6)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Apply selected template
            if template_name in self.templates:
                doc = self.templates[template_name](doc, data)
            else:
                doc = self.build_modern_pro_template(doc, data)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            print(f"Error generating resume: {str(e)}")
            print(traceback.format_exc())
            raise

    def _add_horizontal_line(self, doc, color=RGBColor(200, 200, 200)):
        """Add a horizontal line separator"""
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '%02x%02x%02x' % (color[0], color[1], color[2]))
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)

    def _format_list_items(self, items):
        """Convert items to list"""
        if isinstance(items, str):
            return [item.strip() for item in items.split('\n') if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []

    def _add_section_heading(self, parent, text, color):
        heading = parent.add_paragraph()
        run = heading.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = color
        run.font.name = 'Calibri'
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        return heading

    def _add_text_paragraph(self, parent, text, size=10, bold=False, italic=False, color=None, left_indent=0, space_after=Pt(4)):
        paragraph = parent.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = color
        paragraph.paragraph_format.left_indent = Inches(left_indent)
        paragraph.paragraph_format.space_after = space_after
        return paragraph

    def build_modern_pro_template(self, doc, data):
        """Modern professional template with clean design"""
        try:
            primary = self.color_schemes["Modern Pro"]["primary"]
            secondary = self.color_schemes["Modern Pro"]["secondary"]

            header = doc.add_table(rows=1, cols=2)
            header.autofit = False
            header.columns[0].width = Inches(4.75)
            header.columns[1].width = Inches(2.5)

            name_cell = header.cell(0, 0)
            contact_cell = header.cell(0, 1)

            name_para = name_cell.paragraphs[0]
            name_run = name_para.add_run(data['personal_info'].get('full_name', '').upper())
            name_run.font.size = Pt(28)
            name_run.font.bold = True
            name_run.font.color.rgb = primary
            name_run.font.name = 'Calibri'
            name_para.paragraph_format.space_after = Pt(2)

            if data['personal_info'].get('title'):
                title_para = name_cell.add_paragraph()
                title_run = title_para.add_run(data['personal_info']['title'])
                title_run.font.size = Pt(13)
                title_run.font.color.rgb = secondary
                title_run.font.name = 'Calibri'
                title_para.paragraph_format.space_after = Pt(2)

            contact_para = contact_cell.paragraphs[0]
            contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            contact_items = []
            if data['personal_info'].get('email'):
                contact_items.append(data['personal_info']['email'])
            if data['personal_info'].get('phone'):
                contact_items.append(data['personal_info']['phone'])
            if data['personal_info'].get('location'):
                contact_items.append(data['personal_info']['location'])
            if contact_items:
                contact_run = contact_para.add_run(' | '.join(contact_items))
                contact_run.font.size = Pt(9)
                contact_run.font.color.rgb = RGBColor(80, 80, 80)
                contact_run.font.name = 'Calibri'
                contact_para.paragraph_format.space_after = Pt(3)

            link_items = []
            if data['personal_info'].get('linkedin'):
                link_items.append(f"LinkedIn: {data['personal_info']['linkedin']}")
            if data['personal_info'].get('portfolio'):
                link_items.append(f"Portfolio: {data['personal_info']['portfolio']}")
            if link_items:
                links_para = contact_cell.add_paragraph()
                links_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                links_run = links_para.add_run(' | '.join(link_items))
                links_run.font.size = Pt(9)
                links_run.font.color.rgb = primary
                links_run.font.name = 'Calibri'
                links_para.paragraph_format.space_after = Pt(3)

            self._add_horizontal_line(doc, primary)

            body = doc.add_table(rows=1, cols=2)
            body.autofit = False
            body.columns[0].width = Inches(2.15)
            body.columns[1].width = Inches(5.25)
            left_cell = body.cell(0, 0)
            right_cell = body.cell(0, 1)

            # Sidebar content
            if data.get('skills'):
                self._add_section_heading(left_cell, 'SKILLS', primary)
                skill_categories = [
                    ('technical', 'Technical Skills'),
                    ('soft', 'Soft Skills'),
                    ('languages', 'Languages'),
                    ('tools', 'Tools & Platforms')
                ]
                for key, label in skill_categories:
                    items = self._format_list_items(data['skills'].get(key, []))
                    if items:
                        self._add_text_paragraph(left_cell, f"{label}:", size=10, bold=True, color=secondary, space_after=Pt(0))
                        skills_para = left_cell.add_paragraph(' • '.join(items))
                        skills_para.runs[0].font.size = Pt(9)
                        skills_para.runs[0].font.color.rgb = RGBColor(50, 50, 50)
                        skills_para.runs[0].font.name = 'Calibri'
                        skills_para.paragraph_format.space_after = Pt(4)

            if data.get('education') and len(data['education']) > 0:
                self._add_section_heading(left_cell, 'EDUCATION', primary)
                for edu in data['education']:
                    school_para = left_cell.add_paragraph()
                    school_run = school_para.add_run(edu.get('school', ''))
                    school_run.font.size = Pt(10)
                    school_run.font.bold = True
                    school_run.font.color.rgb = RGBColor(0, 0, 0)
                    school_run.font.name = 'Calibri'
                    school_para.paragraph_format.space_after = Pt(1)

                    degree_text = edu.get('degree', '')
                    if edu.get('field'):
                        degree_text = f"{degree_text} in {edu.get('field')}" if degree_text else f"{edu.get('field')}"
                    if degree_text:
                        degree_para = left_cell.add_paragraph(degree_text)
                        degree_para.runs[0].font.size = Pt(9)
                        degree_para.runs[0].font.color.rgb = secondary
                        degree_para.runs[0].font.name = 'Calibri'
                        degree_para.paragraph_format.space_after = Pt(1)

                    info_text = edu.get('graduation_date', '')
                    if edu.get('gpa'):
                        info_text += f" | GPA: {edu.get('gpa')}"
                    if info_text.strip():
                        info_para = left_cell.add_paragraph(info_text)
                        info_para.runs[0].font.size = Pt(9)
                        info_para.runs[0].font.italic = True
                        info_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                        info_para.runs[0].font.name = 'Calibri'
                        info_para.paragraph_format.space_after = Pt(4)

            # Main body
            if data.get('summary') and data['summary'].strip():
                self._add_section_heading(right_cell, 'PROFESSIONAL SUMMARY', primary)
                summary_para = right_cell.add_paragraph(data['summary'])
                summary_para.runs[0].font.size = Pt(10)
                summary_para.runs[0].font.color.rgb = RGBColor(50, 50, 50)
                summary_para.runs[0].font.name = 'Calibri'
                summary_para.paragraph_format.space_after = Pt(10)

            if data.get('experience') and len(data['experience']) > 0:
                self._add_section_heading(right_cell, 'PROFESSIONAL EXPERIENCE', primary)
                for exp in data['experience']:
                    job_title = right_cell.add_paragraph()
                    position_run = job_title.add_run(exp.get('position', ''))
                    position_run.font.size = Pt(11)
                    position_run.font.bold = True
                    position_run.font.color.rgb = RGBColor(0, 0, 0)
                    position_run.font.name = 'Calibri'

                    company_run = job_title.add_run(f" | {exp.get('company', '')}")
                    company_run.font.size = Pt(10)
                    company_run.font.color.rgb = secondary
                    company_run.font.name = 'Calibri'
                    job_title.paragraph_format.space_after = Pt(2)

                    dates_para = right_cell.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
                    dates_para.runs[0].font.size = Pt(9)
                    dates_para.runs[0].font.italic = True
                    dates_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)
                    dates_para.runs[0].font.name = 'Calibri'
                    dates_para.paragraph_format.space_after = Pt(2)

                    if exp.get('description'):
                        desc_para = right_cell.add_paragraph(exp.get('description'))
                        desc_para.runs[0].font.size = Pt(10)
                        desc_para.runs[0].font.color.rgb = RGBColor(50, 50, 50)
                        desc_para.runs[0].font.name = 'Calibri'
                        desc_para.paragraph_format.space_after = Pt(4)

                    all_items = self._format_list_items(exp.get('responsibilities', []))
                    all_items.extend(self._format_list_items(exp.get('achievements', [])))
                    for item in all_items[:5]:
                        bullet = right_cell.add_paragraph(item, style='List Bullet')
                        bullet.paragraph_format.left_indent = Inches(0.2)
                        bullet.paragraph_format.space_after = Pt(2)
                        bullet.runs[0].font.size = Pt(10)
                        bullet.runs[0].font.name = 'Calibri'
                    right_cell.add_paragraph().paragraph_format.space_after = Pt(6)

            if data.get('projects') and len(data['projects']) > 0:
                self._add_section_heading(right_cell, 'KEY PROJECTS', primary)
                for proj in data['projects'][:3]:
                    proj_name = right_cell.add_paragraph()
                    proj_run = proj_name.add_run(proj.get('name', ''))
                    proj_run.font.size = Pt(11)
                    proj_run.font.bold = True
                    proj_run.font.color.rgb = RGBColor(0, 0, 0)
                    proj_run.font.name = 'Calibri'
                    proj_name.paragraph_format.space_after = Pt(1)

                    if proj.get('technologies'):
                        tech_para = right_cell.add_paragraph(f"Tech: {proj.get('technologies')}")
                        tech_para.runs[0].font.size = Pt(9)
                        tech_para.runs[0].font.italic = True
                        tech_para.runs[0].font.color.rgb = secondary
                        tech_para.runs[0].font.name = 'Calibri'
                        tech_para.paragraph_format.space_after = Pt(2)

                    if proj.get('description'):
                        proj_desc_para = right_cell.add_paragraph(proj.get('description'))
                        proj_desc_para.runs[0].font.size = Pt(10)
                        proj_desc_para.runs[0].font.color.rgb = RGBColor(50, 50, 50)
                        proj_desc_para.runs[0].font.name = 'Calibri'
                        proj_desc_para.paragraph_format.space_after = Pt(4)

                    proj_items = self._format_list_items(proj.get('achievements', []))
                    for item in proj_items[:3]:
                        bullet = right_cell.add_paragraph(item, style='List Bullet')
                        bullet.paragraph_format.left_indent = Inches(0.2)
                        bullet.paragraph_format.space_after = Pt(2)
                        bullet.runs[0].font.size = Pt(10)
                        bullet.runs[0].font.name = 'Calibri'
                    right_cell.add_paragraph().paragraph_format.space_after = Pt(6)

            return doc

        except Exception as e:
            print(f"Error building modern pro template: {str(e)}")
            print(traceback.format_exc())
            raise

    def build_executive_template(self, doc, data):
        """Executive/formal template with traditional styling"""
        return self.build_modern_pro_template(doc, data, template_name="Executive")

    def build_minimalist_template(self, doc, data):
        """Clean minimalist template"""
        return self.build_modern_pro_template(doc, data, template_name="Minimalist")

    def build_corporate_template(self, doc, data):
        """Corporate template with blue tones"""
        return self.build_modern_pro_template(doc, data, template_name="Corporate")

    def build_creative_tech_template(self, doc, data):
        """Creative tech template with modern colors"""
        return self.build_modern_pro_template(doc, data, template_name="Creative Tech")
